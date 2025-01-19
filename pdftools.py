import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
from docx.dml.color import ColorFormat
import secrets
import os.path


#import pypandoc
import subprocess
from werkzeug.utils import secure_filename

HEADING = "听写 Tested Words"
SUBTITLE = "For unknown words, please fill in with 'N/A'.\n"
TABLE_STYLE = "Light Grid Accent 2"
MAIN_FONT = "Arial"
CHINESE_FONT = "SimSun"

ROWHEIGHT = Inches(0.5)
MARGINS = Inches(0.75) # Don't catch you slippin now

TITLE_COLOR = RGBColor(0, 0, 0)

def random_characters(length):
    pca = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890"
    # use python secrets module for security: https://docs.python.org/3/library/secrets.html#module-secrets
    return "".join([secrets.choice(pca) for x in range(0, length - 1)]) 

def _make_tempfilename(file_extension, leng=32):
    filename = random_characters(leng)

    if os.path.exists(filename + file_extension):
        return _make_tempfilename(file_extension)
    else:
        return secure_filename(filename + file_extension)

     


def create_pdf(wordslist : list, output_filename : str):

    print(f"making pdf, filename {output_filename}")

    # muahahahaha type hinting in python lol
    # to accomodate pinyin (or lack thereof), words list can be in format
    '''
    [
    [heading, heading, heading],
    [chinese, english, pinyin],
    [chinese, english, pinyin]
    ]
    '''
    # output filename shouldnt be too important cuz its gonna be temporary anyway.

    newdoc = docx.Document()
    newdoc.sections[0].left_margin = MARGINS
    newdoc.sections[0].right_margin = MARGINS
    newdoc.sections[0].top_margin = MARGINS
    newdoc.sections[0].bottom_margin = MARGINS




    newparagraph = newdoc.add_paragraph(HEADING, style="Heading 1")
    newparagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    newparagraph.runs[0].font.size = Pt(36)
    newparagraph.runs[0].font.color.rgb = TITLE_COLOR

    anewparagraph = newdoc.add_paragraph(SUBTITLE, style="Heading 3")
    anewparagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anewparagraph.runs[0].font.size = Pt(16)
    anewparagraph.runs[0].font.color.rgb = TITLE_COLOR



    table = newdoc.add_table(rows=1, cols=len(wordslist[0]))
    table.style = newdoc.styles[TABLE_STYLE]
    heading_row = table.rows[0].cells

    for i in range(0, len(heading_row)):
            paragraph = heading_row[i].paragraphs[0]
            run = paragraph.add_run(wordslist[0][i])

            run.font.size = Pt(18)
            run.font.name = MAIN_FONT
            run.bold = True

            run._element.rPr.rFonts.set(qn('w:eastAsia'), CHINESE_FONT)

    del wordslist[0]

    for row in wordslist:
        newrow = table.add_row() # 牛肉 lol
        newrow.height = ROWHEIGHT
        row_cells = newrow.cells
        for i in range(0, len(row_cells)):
            paragraph = row_cells[i].paragraphs[0]
            print(row[i])
            run = paragraph.add_run(row[i])
            

            run.font.size = Pt(16)
            run.font.name = MAIN_FONT
            run.bold = False

            run._element.rPr.rFonts.set(qn('w:eastAsia'), CHINESE_FONT)

    newparagraph = newdoc.add_paragraph("Made with Instant Tingxie", style="Normal")
    newparagraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    newparagraph.runs[0].font.size = Pt(8)
    # footer = newdoc.sections[0].footer
    # footer.text = "Made with InstantTingxie"
    tempname = _make_tempfilename(".docx")
    newdoc.save(tempname)

    # output = pypandoc.convert_file(
    #      tempname, 
    #      'pdf', 
    #      outputfile=output_filename
    #      )
    command = ["rocketpdf", "parsedoc", os.getcwd() + "\\" +  tempname, "-o", output_filename]

    #output = subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    subprocess.run(command)
    os.remove(tempname)
    print("created pdf")
    #return output


        


if __name__ == "__main__":
    # just to test it

    demo_wordslist = [
        ['中文', 'English'],
        ['妈妈', 'Mother'],
        ['爸爸', 'Father'],
        ['弟弟', 'Brother'],
    ]
    tingxie_wordslist = [
         ["English", "中文"],
         ["Mother", "妈妈"],
         ["Father", "爸爸"],
         ["Elder brother", "哥哥"],
         ["Elder sister", "姐姐"],
         ["Grandmother", "奶奶"],
         ["Grandfather", "爷爷"],
         ["Child", "孩子"],

    ]
    tingxie_wordslist_real = [
         ["English", "中文"],
         ["Mother", ""],
         ["Father", ""],
         ["Elder brother", ""],
         ["Elder sister", ""],
         ["Grandmother", ""],
         ["Grandfather", ""],
         ["Child", ""],

    ]
    print(create_pdf(tingxie_wordslist_real, "test.pdf"))
    #print(_make_tempfilename(".pdf"))
    





